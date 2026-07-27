from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from document.domain.errors import CorruptDocumentError
from document.domain.models import DocumentFormat
from document.domain.ports import ExtractedText, TextExtractorPort


class PdfExtractor(TextExtractorPort):
    """Text-layer extraction only. Scanned PDFs surface as no extractable text."""

    @property
    def document_format(self) -> DocumentFormat:
        return DocumentFormat.PDF

    def extract(self, payload: bytes) -> ExtractedText:
        try:
            reader = PdfReader(BytesIO(payload))
        except (PdfReadError, OSError, ValueError) as error:
            raise CorruptDocumentError(DocumentFormat.PDF, str(error)) from error

        if reader.is_encrypted:
            try:
                if reader.decrypt("") == 0:
                    raise CorruptDocumentError(
                        DocumentFormat.PDF, "document is password-protected"
                    )
            except (PdfReadError, NotImplementedError) as error:
                raise CorruptDocumentError(DocumentFormat.PDF, str(error)) from error

        warnings: list[str] = []
        pages: list[str] = []
        for index, page in enumerate(reader.pages):
            try:
                pages.append(page.extract_text() or "")
            except (PdfReadError, ValueError, KeyError):
                warnings.append(f"Page {index + 1} could not be read and was skipped.")

        title, authors = self._metadata(reader)
        return ExtractedText(
            text="\n\n".join(page.strip() for page in pages if page.strip()),
            title=title,
            authors=authors,
            page_count=len(reader.pages),
            warnings=warnings,
        )

    def _metadata(self, reader: PdfReader) -> tuple[str | None, list[str]]:
        try:
            metadata = reader.metadata
        except (PdfReadError, ValueError):
            return None, []
        if metadata is None:
            return None, []
        title = metadata.title.strip() if metadata.title else None
        author = metadata.author.strip() if metadata.author else None
        authors = [part.strip() for part in author.split(";")] if author else []
        return title or None, [a for a in authors if a]


class DocxExtractor(TextExtractorPort):
    @property
    def document_format(self) -> DocumentFormat:
        return DocumentFormat.DOCX

    def extract(self, payload: bytes) -> ExtractedText:
        try:
            document = DocxDocument(BytesIO(payload))
        except (PackageNotFoundError, KeyError, ValueError, OSError) as error:
            raise CorruptDocumentError(DocumentFormat.DOCX, str(error)) from error

        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            blocks.append(text)

        blocks.extend(self._table_blocks(document.tables))

        core = document.core_properties
        title = core.title.strip() if core.title else None
        author = core.author.strip() if core.author else None

        return ExtractedText(
            text="\n\n".join(blocks),
            title=title or None,
            authors=[author] if author else [],
        )

    def _table_blocks(self, tables: list[Table]) -> list[str]:
        blocks: list[str] = []
        for table in tables:
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                for row in table.rows
            ]
            populated = [row for row in rows if row]
            if populated:
                blocks.append("\n".join(populated))
        return blocks
